#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成高考志愿智能推荐系统 v6.0 技术报告（Word 文档）
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '4472C4')

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'D9E2F3')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def build_report():
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ===== 封面 =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('高考志愿智能推荐系统 v6.0')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('三层递进漏斗引擎 · 技术架构与算法设计报告')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'文档版本：v6.0  |  生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    doc.add_page_break()

    # ===== 目录 =====
    doc.add_heading('目  录', level=1)
    toc_items = [
        '一、系统概述',
        '二、前端交互设计',
        '三、算法架构总览',
        '四、Layer 1：学科门类初筛',
        '五、Layer 2：专业类精选',
        '六、Layer 3：专业微观狙击',
        '七、特殊赛道机制',
        '八、边界条件与容错设计',
        '九、算法核心优势',
        '十、数据规格',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 一、系统概述 =====
    doc.add_heading('一、系统概述', level=1)

    doc.add_paragraph(
        '高考志愿智能推荐系统 v6.0 是一款面向新高考改革背景下的志愿填报辅助工具。'
        '系统采用「隐蔽式深度测评 + 三层递进漏斗」架构，通过 40 道心理学场景题采集用户'
        '的认知风格、产业向往和微观行为模式，结合用户选科、分数位次、家庭资源和体检限制，'
        '在 883 个本科专业中逐层筛选出最具适配度的推荐结果。'
    )

    doc.add_heading('核心设计理念', level=2)
    principles = [
        ('隐蔽式测评',
         '所有题目面向 18 岁高中生日常场景设计，不使用专业术语。用户无法通过"猜答案"操纵结果，保障信号的真实性。'),
        ('三层递进漏斗',
         '门类初筛 → 专业类精选 → 微观狙击，每一层解决不同粒度的问题，避免扁平化排序的信息丢失。'),
        ('匹配分 ≠ 录取概率',
         '系统只计算「人-专业」的内在适配度，不掺杂院校录取概率。录取风险通过前端标签和文案提示，不污染推荐纯度。'),
        ('选科精确校验',
         '在专业类级别（非学科门类级别）校验新高考选科要求，精确区分「物化双锁」「单物理」「单化学」等不同要求。'),
    ]
    for t, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(f'▸ {t}：')
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(desc)

    # ===== 二、前端交互设计 =====
    doc.add_heading('二、前端交互设计', level=1)

    doc.add_heading('2.1 技术栈', level=2)
    doc.add_paragraph('前端框架：Streamlit（Python 原生 Web 框架）')
    doc.add_paragraph('数据后端：Supabase 云端数据库（卡密管理，防休眠丢数据）')
    doc.add_paragraph('可视化：Plotly 交互式图表')
    doc.add_paragraph('样式：自定义 CSS 主题（Inter 字体 + 紫色渐变配色）')

    doc.add_heading('2.2 页面流程', level=2)
    doc.add_paragraph('系统采用线性引导式流程，用户无法跳过步骤：')
    steps = [
        '① 激活码解锁 —— Supabase 云端校验，一码一人，使用后核销',
        '② 基本信息采集 —— 选考科目、预估分数/位次、家庭经济/城市/资源、体检限制、特殊赛道意向',
        '③ 隐蔽式测评（40题）—— 分 4 块逐步展示，每块一个表单提交，支持断点续答',
        '④ 结果展示 —— 行为画像雷达图 + 三层漏斗摘要 + 专业推荐卡片（可展开）',
        '⑤ 重新测评 —— 一键回到激活码输入页，所有状态清零',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('2.3 管理员后台', level=2)
    doc.add_paragraph(
        '通过侧边栏密码解锁管理员面板，提供：批量生成激活码（密码学安全随机，XXXX-XXXX-XXXX 格式）、'
        '卡密看板（总生成数/已使用/未使用/转化率）、CSV 导出功能。'
    )

    doc.add_heading('2.4 视觉设计', level=2)
    doc.add_paragraph('• 渐变色卡片（#667EEA → #764BA2）展示推荐结果，金牌/银牌/铜牌标记前三名')
    doc.add_paragraph('• Plotly 柱状图展示 10 维微观行为画像和 10 大产业向往')
    doc.add_paragraph('• 可展开卡片（st.expander）展示每个专业类的详细推荐理由和专业卡片网格')
    doc.add_paragraph('• 响应式布局：st.columns 实现多列网格，适配宽屏')

    # ===== 三、算法架构总览 =====
    doc.add_heading('三、算法架构总览', level=1)

    doc.add_paragraph(
        '核心算法由三层递进漏斗（Funnel Engine）构成，每层解决不同粒度的匹配问题。'
        '漏斗结构确保候选集逐层收窄：883 专业 → 93 专业类 → 8 专业类 → ≤48 专业。'
    )

    doc.add_paragraph(
        '数据流：用户画像（UserProfile）携带硬约束（选科、体检）、软信号（10维微观行为向量、'
        '10维产业向往向量、RIASEC人格向量）和家庭资源信息，依次通过三层漏斗，最终输出结构化推荐结果。'
    )

    add_styled_table(doc,
        ['层级', '名称', '输入', '机制', '输出'],
        [
            ['Layer 1', '学科门类初筛', '认知风格 + 人格倾向',
             '余弦相似度 × 权重', '13 门类得分排序'],
            ['Layer 2', '专业类精选', 'L1得分 + 产业 + 资产 + 分数 + 选科',
             'Jaccard + 矩阵查表 + 硬校验 + Top 8', '≤8 专业类'],
            ['Layer 3', '专业微观狙击', 'L2类别 + 微观动作 + 热度 + 红线',
             '余弦 + 矩阵查表 + 开关 × 市场容量', '每类 ≤6 专业'],
        ],
        col_widths=[1.8, 2.8, 4.0, 3.5, 2.8],
    )

    doc.add_page_break()

    # ===== 四、Layer 1 =====
    doc.add_heading('四、Layer 1：学科门类初筛', level=1)

    doc.add_heading('4.1 设计目标', level=2)
    doc.add_paragraph(
        '在 13 个学科门类（哲学、经济学、法学、教育学、文学、历史学、理学、工学、农学、医学、'
        '军事学、管理学、艺术学）中，基于用户的认知风格和人格倾向进行初筛，'
        '为 Layer 2 提供学科门类级别的先验得分。'
    )

    doc.add_heading('4.2 数学模型', level=2)
    doc.add_paragraph('对于每个学科门类 d，计算：')
    p = doc.add_paragraph()
    run = p.add_run('  Score(d) = Cosine(User_Behavior, Disc_Cognitive) × 0.5')
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run('          + Cosine(User_Personality, Disc_Persona) × 0.3')
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run('          + Discipline_Weight × 0.2')
    run.font.size = Pt(10)

    doc.add_paragraph(
        '权重设计逻辑：认知风格在学科门类层面最具区分度（如逻辑推演 → 理学，共情表达 → 教育学），'
        '赋予 50% 最高权重；人格倾向提供大类匹配（如研究型 I → 理学/医学），赋予 30%；'
        '门类权重来自数据标注（考虑学科体量和覆盖面），赋予 20%。'
    )

    doc.add_heading('4.3 标签向量化', level=2)
    doc.add_paragraph(
        '每个学科门类的「认知风格标签」和「人格倾向标签」通过映射表转换为数值向量。'
        '例如，学科门类「工学」的认知风格标签为 ["系统建构", "逻辑推演", "实验思维"]，'
        '通过 COGNITIVE_STYLE_TO_BEHAVIOR 映射为 10 维行为向量，'
        '再与用户的 10 维微观行为向量做余弦相似度计算。'
    )
    doc.add_paragraph(
        '余弦相似度的设计选择：归一化了向量长度——即使用户所有维度都打高分，不影响匹配方向。'
        '有效防止"永远选 A"的作弊行为，只看用户行为模式的相对强弱。'
    )

    doc.add_heading('4.4 设计决策：选科不参与 L1', level=2)
    doc.add_paragraph(
        'Layer 1 不进行选科过滤。这一决策的原因是：新高考选科要求在专业类级别存在显著差异——'
        '同一学科门类下，不同专业类的选科要求可能完全不同。例如，工学门类中，计算机类只需物理，'
        '而材料类需要物理+化学。在门类级别施加统一的选科系数会导致"一刀切"误判。'
        '精确的选科校验下沉到 Layer 2（专业类级别）。'
    )

    # ===== 五、Layer 2 =====
    doc.add_heading('五、Layer 2：专业类精选', level=1)

    doc.add_heading('5.1 设计目标', level=2)
    doc.add_paragraph(
        '在 93 个专业类中筛选出 Top 8，综合考虑产业向往、家庭资产匹配、分数位次匹配、'
        '新高考选科要求、特殊赛道意向和 L1 门类得分传导。'
    )

    doc.add_heading('5.2 处理流程（7步）', level=2)
    steps_l2 = [
        'Step 1: 选科硬校验 —— 不满足 CATEGORY_SUBJECT_REQUIREMENTS → 直接跳过（不进入计分）',
        'Step 2: 产业匹配 —— Jaccard 指数（用户顶层产业 ∩ 专业类产业标签）',
        'Step 3: 特殊赛道产业拉升 —— 强烈意向时，强制 industry_match ≥ 地板值',
        'Step 4: 资产匹配 —— 家庭经济水平 vs 专业类资产敏感度，查表得分',
        'Step 5: 分数匹配 —— 排名位次 vs 专业类分数敏感度，查表得分',
        'Step 6: 综合计分 —— 产业×0.5 + 资产×0.2 + 分数×0.3，叠加产业热度微调、选科加成、赛道提权、L1传导',
        'Step 7: Top 8 截断 + 多样性强制打散',
    ]
    for s in steps_l2:
        doc.add_paragraph(s, style='List Number')

    doc.add_heading('5.3 选科硬校验（核心创新）', level=2)
    doc.add_paragraph(
        '系统内置 66 个专业类的精确选科要求字典（CATEGORY_SUBJECT_REQUIREMENTS），'
        '数据来源于教育部《普通高校本科招生专业选考科目要求指引》。'
    )

    add_styled_table(doc,
        ['要求类型', '示例专业类', '必选科目', '用户不满足时'],
        [
            ['物化双锁', '临床医学类、口腔医学类、材料类、化学类', '物理 + 化学', '直接跳过'],
            ['单物理', '计算机类、电子信息类、机械类、自动化类', '物理', '直接跳过'],
            ['单化学', '药学类、中药学类、植物生产类', '化学', '直接跳过'],
            ['单地理', '地理科学类', '地理', '直接跳过'],
            ['不限', '工商管理类、法学类、中国语言文学类等', '无', '正常计分'],
        ],
        col_widths=[2.0, 5.0, 2.8, 2.0],
    )

    doc.add_paragraph(
        '精确定义意味着：「物生地」组合（缺化学）无法看到临床医学、口腔医学、材料类等物化双锁类别，'
        '但可以看到计算机类、电子信息类等仅需物理的类别。这与真实高考志愿填报规则完全一致。'
    )

    doc.add_heading('5.4 产业匹配（Jaccard 指数）', level=2)
    doc.add_paragraph(
        '定义用户的「顶层产业」为宏观产业向量中得分 >60 的集群。'
        '对于专业类 c，计算用户顶层产业集合 U 与 c 的产业标签集合 C 的 Jaccard 指数：'
    )
    p = doc.add_paragraph()
    run = p.add_run('  industry_match = |U ∩ C| / |U ∪ C|')
    run.font.size = Pt(10)
    doc.add_paragraph(
        'Jaccard 指数同时对「用户感兴趣但类别不提供」和「类别提供但用户不感兴趣」两方向进行惩罚，'
        '比简单的交集/用户集合（Recall）或交集/类别集合（Precision）更公平。'
        '对于多标签覆盖的专业类，附赠 5%/标签的广度微调（最多 40%），体现跨学科类别的职业灵活性。'
    )

    doc.add_heading('5.5 资产与分数匹配（矩阵查表）', level=2)
    doc.add_paragraph(
        '资产敏感度匹配和分数敏感度匹配均采用矩阵查表方式，而非连续函数。'
        '原因：社会经济数据和高考位次数据天然具有分段特征，离散化匹配更稳健，'
        '避免了连续函数对极端值的过度敏感。'
    )
    doc.add_paragraph(
        '资产匹配矩阵的行（用户经济水平：高/中/低）× 列（专业敏感度：高/中/低），'
        '体现「高经济水平→高敏感也能驾驭」「低经济水平→低敏感最友好」的经济理性。'
    )
    doc.add_paragraph(
        '分数匹配矩阵的行（用户位次段：top/upper/middle/lower）× 列（专业敏感度：极高/高/中/低），'
        '体现「前10%→极高敏感也能打」「后40%→低敏感最友好」的竞争现实。'
    )

    doc.add_heading('5.6 选科匹配加成', level=2)
    doc.add_paragraph(
        '满足该专业类必选科目的，每个科目给予 +3% 微量加成。例如：临床医学类需物理+化学，'
        '满足则得分 ×1.06。这是一个微小但持续的偏向，在不影响主要排序的前提下，'
        '对满足选科要求的专业类提供正向激励。'
    )

    doc.add_heading('5.7 L1 门类得分传导', level=2)
    doc.add_paragraph('Layer 1 的学科门类得分通过传导公式渗入 Layer 2：')
    p = doc.add_paragraph()
    run = p.add_run('  L2_score = L2_score × (0.75 + 0.50 × L1_Discipline_Score)')
    run.font.size = Pt(10)
    doc.add_paragraph(
        '该公式意味着 L1 得分 1.0 的门类获得 +25% 传导加成，L1 得分 0 的门类承受 -25% 抑制。'
        '±25% 的范围确保认知/人格匹配能影响专业类排序，但不主导最终结果（主导权仍在 L2 自身的多维匹配）。'
    )

    doc.add_page_break()

    # ===== 六、Layer 3 =====
    doc.add_heading('六、Layer 3：专业微观狙击', level=1)

    doc.add_heading('6.1 设计目标', level=2)
    doc.add_paragraph(
        '在 Layer 2 选出的 Top 8 专业类中，对每个类别下的具体专业进行微观级别的匹配计算，'
        '每类最多输出 6 个专业。'
    )

    doc.add_heading('6.2 数学模型', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '  Major_Score = (Micro_Match × 0.6 + Heat_Align × 0.4) '
        '× Enrollment_Capacity × Threshold_Pass'
    )
    run.font.size = Pt(10)

    doc.add_paragraph('其中：')
    doc.add_paragraph('• Micro_Match：用户 10 维微观行为向量与专业微观动作标签的余弦相似度（权重 60%）')
    doc.add_paragraph('• Heat_Align：用户风险容忍度与专业社会热度（极高/高/中/低）的矩阵匹配分（权重 40%）')
    doc.add_paragraph('• Enrollment_Capacity：招生体量市场容量系数（0.70 ~ 1.15），体量越大越容易进入')
    doc.add_paragraph('• Threshold_Pass：红线开关，0 或 1。体检限制命中 → 直接归零')

    doc.add_heading('6.3 微观动作匹配', level=2)
    doc.add_paragraph(
        '这是整个漏斗中最底层的信号，也是权重最高的单维度（60%）。'
        '每个专业标注了 1-4 个微观动作标签（如临床医学标注「临床诊断与鉴别」「手术/介入操作」'
        '「医学影像判读」），通过 MICRO_ACTION_TO_BEHAVIOR 映射表转换为 10 维行为向量，'
        '再与用户在 30 道微观场景题中展现的行为模式做余弦匹配。'
    )
    doc.add_paragraph(
        '设计理由：微观行为是最难伪造的信号——一个学生可以在问卷中说"我想学医"，'
        '但无法伪装自己在「精细操作」「记忆积累」「抗压能力」等维度上的真实行为倾向。'
        '这些维度正是通过 30 道日常场景题（实验课、小组作业、B站沉迷内容等）间接测量的。'
    )

    doc.add_heading('6.4 红线机制', level=2)
    doc.add_paragraph(
        'Layer 3 的硬红线（Threshold）采用语义扩展匹配。用户的体检限制标签（如"色盲"）'
        '通过 THRESHOLD_SEMANTIC_MAP 扩展为同义词集合（{"色盲", "红绿色盲", "色觉", "色弱"}），'
        '与专业的红线标签做子串匹配。命中任何一个 → Threshold_Pass = 0 → 该专业直接归零，不可恢复。'
    )
    doc.add_paragraph(
        '当多个 Top 8 类别因红线被大规模清空时，系统自动触发「救援机制」：'
        '从 Layer 2 的候选池（第 9 名起）中提取未参与 L3 的类别，运行 Layer 3 并补充到输出中。'
        '救援阈值为：输出类别 ≤5 个 或 总专业数 <15 个。'
    )

    doc.add_page_break()

    # ===== 七、特殊赛道 =====
    doc.add_heading('七、特殊赛道机制', level=1)

    doc.add_heading('7.1 赛道定义', level=2)
    doc.add_paragraph(
        '系统支持三条特殊赛道：医学、师范、军警。用户在基本信息采集中可选择赛道和态度'
        '（强烈意向 / 可以接受 / 极度抗拒）。'
    )

    add_styled_table(doc,
        ['赛道', '涵盖专业类', '强烈意向', '可以接受', '极度抗拒', '未选择'],
        [
            ['医学', '临床/口腔/基础/中医/中西医/\n药学/中药学/法医/医学技术/\n护理/公卫预防（11类）',
             '产业拉升≥0.65\n分数×1.50', '正常计算', '全部过滤', '正常计算'],
            ['师范', '教育学类、体育学类\n（2类）',
             '产业拉升≥0.60\n分数×1.30', '正常计算', '全部过滤', '正常计算'],
            ['军警', '公安学类、公安技术类、\n兵器类（3类）',
             '产业拉升≥0.60\n分数×1.50\n+ 体检合格', '全部过滤', '全部过滤', '全部过滤'],
        ],
        col_widths=[1.5, 4.0, 3.0, 2.0, 2.0, 2.0],
    )

    doc.add_heading('7.2 设计逻辑', level=2)
    doc.add_paragraph(
        '赛道的核心设计理念是「用户明确意向 > 问卷推测」。即使用户在 40 道问卷中完全没有'
        '表现出对医学产业的向往（例如全部选择了技术类选项），只要用户在基本信息中明确选择'
        '「医学·强烈意向」，系统就会强制拉升医学类的产业匹配分到 0.65 并获得 1.50 倍的分数加成。'
    )
    doc.add_paragraph(
        '同时，强烈意向赛道会跳过 Layer 2 的多样性强制打散——用户既然明确表达了偏好，'
        '系统应当尊重这一选择，不应该用"多样性"稀释用户的意向方向。'
    )

    doc.add_heading('7.3 军警特殊处理', level=2)
    doc.add_paragraph(
        '军警赛道与其他赛道不同：非「强烈意向」一律过滤（即使用户选了「可以接受」）。'
        '这是因为军警方向涉及特殊的职业承诺和生活方式，不同于普通的专业选择，'
        '系统选择保守策略——宁可漏推，不可误导。'
    )

    # ===== 八、边界条件 =====
    doc.add_heading('八、边界条件与容错设计', level=1)

    add_styled_table(doc,
        ['边界条件', '系统行为', '设计理由'],
        [
            ['未选科目', '仅显示不限选科的专业类（30个）', '严格遵循新高考政策'],
            ['价值观向量为空', '风险容忍度默认50，归类为"中"', '中性默认，不影响排序'],
            ['所有产业分 <60', '默认产业设为"互联网与软件"', '最泛化的产业，覆盖面最广'],
            ['极端位次 (0.5% ~ 100%)', '正常计算，分数匹配矩阵自动适配', '矩阵覆盖全范围'],
            ['多体检条件叠加', '逐条命中即触发红线，救援机制兜底', '安全优先，不遗漏限制'],
            ['空选科 + 空价值观', '正常返回（不崩溃），结果偏管理类', '兜底设计优先于报错'],
            ['L3 红线清空 ≥3 类别', '救援机制从备选池补充', '避免结果页空白'],
            ['多样性 ≥5/8 同门类', '逐步替换末位直到门类≤4', '防止回音壁效应'],
        ],
        col_widths=[4.0, 5.5, 4.5],
    )

    # ===== 九、算法核心优势 =====
    doc.add_heading('九、算法核心优势', level=1)

    advantages = [
        (
            '三层递进漏斗架构',
            '不同于传统的「打分 → 排序」扁平化推荐，三层漏斗在每一层解决不同粒度的匹配问题：'
            '门类（认知/人格匹配）、专业类（产业/资产/分数匹配）、具体专业（微观行为/热度/红线）。'
            '这种架构避免了单一维度主导排序的偏差，确保推荐的多样性和层次感。'
            '每层独立计分、逐层截断，保证候选集从 883 有序收窄到 ≤48 个专业，既不过度也不稀疏。'
        ),
        (
            '隐蔽式深度测评',
            '40 道题目全部采用 18 岁高中生日常场景（实验课表现、B站沉迷内容、小组作业角色等），'
            '被测者无法通过「猜答案」来操纵推荐方向。10 条测谎规则（CONSISTENCY_RULES）检测宏观意愿与微观行为的矛盾，'
            '对「叶公好龙」式回答自动衰减 50%，提升信号真实度。'
            '测评结果不仅输出推荐，还提供 10 维行为画像和 RIASEC 人格推断，帮助用户了解自己的底层能力结构。'
        ),
        (
            '新高考选科精确校验',
            '在专业类级别（93 类）逐一定义选科要求，而非在学科门类级别（13 类）一刀切。'
            '精确区分「物化双锁」（临床/口腔/材料）、「单物理」（计算机/电子/机械）、'
            '「单化学」（药学/农学）等不同要求。这是同类系统中罕见的细粒度处理。'
            '选科要求来源于教育部官方指引，覆盖 66 个专业类，空集合表示不限选科，语义清晰无歧义。'
        ),
        (
            '多维匹配矩阵',
            '产业匹配（Jaccard 指数，双向惩罚）、资产匹配（社会经济分段查表）、分数匹配（位次分段查表）'
            '三套独立的匹配机制并行工作，每套机制都有明确的数学定义和经济/社会逻辑支撑。'
            '矩阵查表替代连续函数，避免了过度拟合和极端值敏感。'
            '权重分配经过多轮验证：产业 50%（用户兴趣主导）、资产 20%（辅助性因素）、分数 30%（竞争现实）。'
        ),
        (
            '特殊赛道覆盖',
            '医学/师范/军警三条赛道独立运转，包含产业拉升、分数提权和阻断逻辑。'
            '「强烈意向」会跳过多样性打散，尊重用户的明确选择。'
            '赛道机制确保用户不必在问卷中「表演」出对某方向的兴趣才能获得相关推荐——'
            '用户的明确声明直接覆盖问卷信号。同时，军警的保守策略（非强烈即拦截）体现了对特殊职业选择的审慎态度。'
        ),
        (
            '容错与兜底机制',
            '多层容错设计确保系统在任何边界条件下都不会崩溃或返回空结果：'
            '默认产业兜底、默认风险容忍度、红线救援机制、多样性打散保护、'
            '空选科自动过滤。所有异常路径都有优雅降级方案，用户体验不会因极端输入而中断。'
        ),
        (
            '匹配分与录取概率分离',
            '核心设计原则：匹配分只衡量「人-专业」的内在适配度（你适不适合这个专业），'
            '不掺杂院校录取概率（你能不能考上这个专业）。录取相关的风险提示通过前端标签'
            '（[高分敏感]、[极度内卷]、[排雷预警]）和文案传递，不污染推荐排序。'
            '这避免了「因为他分数低，所以推荐差专业」的陷阱——分数低的同学同样有权利知道自己最适合什么。'
        ),
    ]

    for title, desc in advantages:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ===== 十、数据规格 =====
    doc.add_heading('十、数据规格', level=1)

    add_styled_table(doc,
        ['数据文件', '内容', '规模'],
        [
            ['layer1_disciplines.json', '13个学科门类的认知风格标签和人格倾向标签', '13 条目'],
            ['layer2_categories.json', '93个专业类的产业标签/资产敏感度/分数敏感度', '93 条目'],
            ['layer3_majors.json', '883个专业的微观动作/红线/热度/招生体量', '883 条目 × 5 属性'],
            ['gaokao_majors.xlsx', '专业元数据（序号/门类/专业类/代码/名称）', '883 行 × 5 列'],
            ['questionnaire.py', '40道测评题（10宏观 + 30微观）+ 10条测谎规则', '约 750 行代码'],
            ['enhanced_majors.json', '旧版引擎专业库（含五维标签），v6.0 不再使用', '883 条目'],
        ],
        col_widths=[4.0, 7.0, 3.5],
    )

    doc.add_paragraph()
    doc.add_heading('数据来源说明', level=2)
    doc.add_paragraph('• 专业目录和学科分类：依据教育部《普通高等学校本科专业目录》')
    doc.add_paragraph('• 新高考选科要求：依据教育部《普通高校本科招生专业选考科目要求指引》')
    doc.add_paragraph('• 微观动作标签和产业标签：由领域专家标注，经 LLM 辅助校验')
    doc.add_paragraph('• 社会热度和招生体量：综合近年高考录取数据和行业趋势分析')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 报告结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    run.italic = True

    output_path = '高考志愿推荐系统v6.0_技术报告.docx'
    doc.save(output_path)
    print(f'Done: {output_path}')
    return output_path


if __name__ == '__main__':
    build_report()
